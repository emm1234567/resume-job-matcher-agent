"""Word(.docx) 解析器，基于 python-docx。

注意：python-docx 不支持旧版 .doc 二进制格式。
若遇到 .doc，建议先用 LibreOffice 等工具转换为 .docx 再处理。
"""
from __future__ import annotations

from docx import Document


class DocxParser:
    def parse(self, path: str) -> str:
        doc = Document(path)
        chunks: list[str] = []

        # 抽取正文段落
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if text:
                chunks.append(text)

        # 抽取表格内容（逐单元格拼接，保留行列结构）
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    chunks.append(" | ".join(cells))

        return "\n".join(chunks)
